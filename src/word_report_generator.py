import json
from docx import Document
import os

def generate_word_report():
    with open("reports/insights.json", "r") as f:
        insights = json.load(f)

    doc = Document()
    doc.add_heading("Power AI - Automated Data Analysis Report", level=1)

    doc.add_heading("Dataset Overview", level=2)
    doc.add_paragraph(f"Total Records: {insights['total_records']}")
    doc.add_paragraph(f"Total Columns: {insights['total_columns']}")

    doc.add_heading("Average Salary by Department", level=2)
    for dept, salary in insights["average_salary_by_department"].items():
        doc.add_paragraph(f"{dept}: {salary}")

    doc.add_heading("Key Insight", level=2)
    doc.add_paragraph(
        f"The highest salary is observed in the "
        f"{insights['highest_salary']['department']} department "
        f"with a value of {insights['highest_salary']['value']}."
    )

    os.makedirs("reports", exist_ok=True)
    output_path = "reports/PowerAI_Report.docx"
    doc.save(output_path)

    return output_path
